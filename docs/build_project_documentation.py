from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT_PATH = Path(__file__).with_name("dokumentacja_projektowa_budgetly_uml.docx")
DIAGRAM_DIR = Path(__file__).with_name("uml_diagrams")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2EC"
DIAGRAM_LINE = "#2E74B5"
DIAGRAM_FILL = "#F4F6F9"
DIAGRAM_ACCENT = "#E8EEF5"

PL_REPLACEMENTS = {
    "Programowanie aplikacji internetowych": "Programowanie aplikacji internetowych",
    "Aplikacja webowa do zarzadzania budzetem osobistym": "Aplikacja webowa do zarządzania budżetem osobistym",
    "Budgetly - wielouzytkownikowy dziennik wydatkow": "Budgetly - wieloużytkownikowy dziennik wydatków",
    "Karol Scislowski": "Karol Ścisłowski",
    "dokumentacja wstepna": "dokumentacja wstępna",
    "Dokumentacja wstepna": "Dokumentacja wstępna",
    "Krotki": "Krótki",
    "krotki": "krótki",
    "krotkiego": "krótkiego",
    "Mierzalne": "Mierzalne",
    "mierzalne": "mierzalne",
    "wartosc": "wartość",
    "Wartosc": "Wartość",
    "wartosci": "wartości",
    "Szczegolna": "Szczególna",
    "szczegolnej": "szczególnej",
    "Szczegolne": "Szczególne",
    "spelnienia": "spełnienia",
    "spelnione": "spełnione",
    "kryteriow": "kryteriów",
    "celow": "celów",
    "dwoch": "dwóch",
    "wybor": "wybór",
    "wyboru": "wyboru",
    "aktorow": "aktorów",
    "obiektow": "obiektów",
    "ramow": "ramowych",
    "Uwzgledniono": "Uwzględniono",
    "uwzgledniono": "uwzględniono",
    "wstepna": "wstępna",
    "Wstepna": "Wstępna",
    "systemow": "systemów",
    "finansow": "finansów",
    "rejestrowac": "rejestrować",
    "ustawiac": "ustawiać",
    "eksportowac": "eksportować",
    "kazdego": "każdego",
    "sa": "są",
    "stalych": "stałych",
    "budzetowe": "budżetowe",
    "budzetowych": "budżetowych",
    "pomagaja": "pomagają",
    "rozpoznac": "rozpoznać",
    "zachowac": "zachować",
    "kopie": "kopię",
    "wykonac": "wykonać",
    "dodatkowa": "dodatkową",
    "analize": "analizę",
    "koncowych": "końcowych",
    "koncowego": "końcowego",
    "Zagrozenia": "Zagrożenia",
    "zagrozenia": "zagrożenia",
    "systemu": "systemu",
    "projektow": "projektów",
    "projektu": "projektu",
    "oceniajacy": "oceniający",
    "oceniania": "oceniania",
    "Prowadzacy": "Prowadzący",
    "kompletnosc": "kompletność",
    "spojnosc": "spójność",
    "spojna": "spójna",
    "spojny": "spójny",
    "spojne": "spójne",
    "spójne": "spójne",
    "wielouzytkownikowa": "wieloużytkownikowa",
    "wielouzytkownikowy": "wieloużytkownikowy",
    "uzytkownikow": "użytkowników",
    "Uzytkownikow": "Użytkowników",
    "uzytkownika": "użytkownika",
    "Uzytkownika": "Użytkownika",
    "uzytkownik": "użytkownik",
    "Uzytkownik": "Użytkownik",
    "uzytkownikiem": "użytkownikiem",
    "uzytkownikowi": "użytkownikowi",
    "Uzytecznosc": "Użyteczność",
    "uzyc": "użyć",
    "uzywalna": "używalna",
    "zalogowany": "zalogowany",
    "zalogowanego": "zalogowanego",
    "zarzadzania": "zarządzania",
    "zarzadzac": "zarządzać",
    "zarzadzaj": "zarządzaj",
    "budzetem": "budżetem",
    "budzetami": "budżetami",
    "budzetu": "budżetu",
    "budzetow": "budżetów",
    "budzety": "budżety",
    "budzet": "budżet",
    "Budzety": "Budżety",
    "Budzet": "Budżet",
    "Budzetowanie": "Budżetowanie",
    "wydatkow": "wydatków",
    "przychodow": "przychodów",
    "kosztow": "kosztów",
    "rachunkow": "rachunków",
    "abonamentow": "abonamentów",
    "transakcje": "transakcje",
    "kategoryzowac": "kategoryzować",
    "automatyzowac": "automatyzować",
    "porownywac": "porównywać",
    "porownuje": "porównuje",
    "porownania": "porównania",
    "porownanie": "porównanie",
    "Porownanie": "Porównanie",
    "porownan": "porównań",
    "miesiecy": "miesięcy",
    "miesiace": "miesiące",
    "miesiacach": "miesiącach",
    "miesiac": "miesiąc",
    "miesiaca": "miesiąca",
    "miesiacami": "miesiącami",
    "miesieczne": "miesięczne",
    "miesieczny": "miesięczny",
    "miesiecznego": "miesięcznego",
    "miesiecznym": "miesięcznym",
    "miesiecznych": "miesięcznych",
    "biezacy": "bieżący",
    "Biezacy": "Bieżący",
    "sytuacje": "sytuację",
    "przekroczen": "przekroczeń",
    "ogolna": "ogólna",
    "ogolny": "ogólny",
    "ogolnego": "ogólnego",
    "ogolnosci": "ogólności",
    "powiazania": "powiązania",
    "powiazanie": "powiązanie",
    "Powiazanie": "Powiązanie",
    "powiazan": "powiązań",
    "Glowne": "Główne",
    "glowne": "główne",
    "glownych": "głównych",
    "glownym": "głównym",
    "wymagan": "wymagań",
    "Zespolu": "Zespołu",
    "zespolu": "zespołu",
    "zespol": "zespół",
    "Zespol": "Zespół",
    "odpowiedzialnosci": "odpowiedzialności",
    "Odpowiedzialnosc": "Odpowiedzialność",
    "odpowiedzialnosc": "odpowiedzialność",
    "czestotliwosc": "częstotliwość",
    "czestotliwosci": "częstotliwości",
    "wystapienia": "wystąpienia",
    "walute": "walutę",
    "dostepu": "dostępu",
    "dostep": "dostęp",
    "dostepna": "dostępna",
    "dostepne": "dostępne",
    "dostarcza": "dostarcza",
    "bledow": "błędów",
    "Bledna": "Błędna",
    "bledne": "błędne",
    "blednych": "błędnych",
    "błedne": "błędne",
    "obsluga": "obsługa",
    "obslugi": "obsługi",
    "obslugiwane": "obsługiwane",
    "obsluguje": "obsługuje",
    "wejscia": "wejścia",
    "przeplywow": "przepływów",
    "przeplyw": "przepływ",
    "wplyw": "wpływ",
    "Wplyw": "Wpływ",
    "zrodlo": "źródło",
    "Zrodlo": "Źródło",
    "srodowiskowe": "środowiskowe",
    "srodowiskowych": "środowiskowych",
    "srodowiska": "środowiska",
    "srodowisko": "środowisko",
    "mozna": "można",
    "moze": "może",
    "Moze": "Może",
    "latwo": "łatwo",
    "ulatwia": "ułatwia",
    "Ulatwia": "Ułatwia",
    "roznic": "różnic",
    "roznice": "różnice",
    "roznica": "różnica",
    "zaleznosc": "zależność",
    "zewnetrznego": "zewnętrznego",
    "zewnetrzny": "zewnętrzny",
    "Zewnetrzny": "Zewnętrzny",
    "wewnetrzny": "wewnętrzny",
    "Wewnetrzny": "Wewnętrzny",
    "bezpieczenstwo": "bezpieczeństwo",
    "Bezpieczenstwo": "Bezpieczeństwo",
    "czytelnosc": "czytelność",
    "wersje": "wersję",
    "demonstracje": "demonstrację",
    "sprawdzic": "sprawdzić",
    "sprawdzaja": "sprawdzają",
    "Jakosc": "Jakość",
    "jakosc": "jakość",
    "jakosci": "jakości",
    "Poprawnosc": "Poprawność",
    "poprawnosc": "poprawność",
    "poprawnosci": "poprawności",
    "wejsciowych": "wejściowych",
    "rownolegle": "równolegle",
    "pracowac": "pracować",
    "uruchomic": "uruchomić",
    "Przechodzacy": "Przechodzący",
    "przechod": "przejście",
    "Najwazniejsze": "Najważniejsze",
    "najwazniejsze": "najważniejsze",
    "funkcjonalnosc": "funkcjonalność",
    "poziomow": "poziomów",
    "odpowiadac": "odpowiadać",
    "udostepnic": "udostępnić",
    "trafia": "trafią",
    "ktorych": "których",
    "ktory": "który",
    "ktora": "która",
    "ktore": "które",
    "ktorym": "którym",
    "ktorych": "których",
    "zostal": "został",
    "zostala": "została",
    "zostaly": "zostały",
    "wlasne": "własne",
    "wlasnych": "własnych",
    "wlasnymi": "własnymi",
    "wlasciwe": "właściwe",
    "pelni": "pełni",
    "pelny": "pełny",
    "pelna": "pełna",
    "pelne": "pełne",
    "swoje": "swoje",
    "srodki": "środki",
}


def polish_text(text: str) -> str:
    for source, target in sorted(PL_REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True):
        text = re.sub(rf"(?<![A-Za-z0-9_]){re.escape(source)}(?![A-Za-z0-9_])", target, text)
    return text


def diagram_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if text_size(draw, candidate, font)[0] <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    fill: str = "#111827",
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 24)
    line_height = text_size(draw, "Ag", font)[1] + 6
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1 - total_height) // 2)
    for line in lines:
        width, _ = text_size(draw, line, font)
        draw.text((x1 + (x2 - x1 - width) // 2, y), line, font=font, fill=fill)
        y += line_height


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = DIAGRAM_LINE,
    width: int = 4,
) -> None:
    import math

    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = 0.45
    p1 = (
        end[0] - length * math.cos(angle - spread),
        end[1] - length * math.sin(angle - spread),
    )
    p2 = (
        end[0] - length * math.cos(angle + spread),
        end[1] - length * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill=fill)


def draw_connector(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
) -> None:
    draw.line([start, end], fill=DIAGRAM_LINE, width=3)
    if label:
        font = diagram_font(22)
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2
        w, h = text_size(draw, label, font)
        draw.rectangle((mx - w // 2 - 6, my - h // 2 - 4, mx + w // 2 + 6, my + h // 2 + 4), fill="white")
        draw.text((mx - w // 2, my - h // 2), label, font=font, fill="#111827")


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    stroke = "#111827"
    font = diagram_font(24, bold=True)
    draw.ellipse((x - 20, y, x + 20, y + 40), outline=stroke, width=4)
    draw.line((x, y + 40, x, y + 115), fill=stroke, width=4)
    draw.line((x - 45, y + 70, x + 45, y + 70), fill=stroke, width=4)
    draw.line((x, y + 115, x - 40, y + 170), fill=stroke, width=4)
    draw.line((x, y + 115, x + 40, y + 170), fill=stroke, width=4)
    draw_centered_text(draw, (x - 95, y + 182, x + 95, y + 240), label, font)


def draw_use_case(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], label: str) -> None:
    draw.ellipse(box, outline=DIAGRAM_LINE, width=4, fill="white")
    draw_centered_text(draw, box, label, diagram_font(24))


def draw_component(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str = "") -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, outline=DIAGRAM_LINE, width=4, fill=DIAGRAM_FILL)
    draw.rectangle((x1 + 18, y1 + 18, x1 + 58, y1 + 38), outline=DIAGRAM_LINE, width=3, fill="white")
    draw.rectangle((x1 + 18, y1 + 46, x1 + 58, y1 + 66), outline=DIAGRAM_LINE, width=3, fill="white")
    draw_centered_text(draw, (x1 + 70, y1 + 10, x2 - 10, y1 + 58), title, diagram_font(27, bold=True))
    if body:
        draw_centered_text(draw, (x1 + 20, y1 + 70, x2 - 20, y2 - 10), body, diagram_font(22), fill="#374151")


def draw_class_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    attributes: list[str],
    methods: list[str],
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=8, outline="#111827", width=3, fill="white")
    title_h = 48
    method_h = max(44, 22 + 24 * len(methods))
    attr_h = max(86, (y2 - y1) - title_h - method_h)
    draw.rectangle((x1, y1, x2, y1 + title_h), outline="#111827", width=3, fill=DIAGRAM_ACCENT)
    draw.line((x1, y1 + title_h + attr_h, x2, y1 + title_h + attr_h), fill="#111827", width=3)
    draw_centered_text(draw, (x1, y1, x2, y1 + title_h), title, diagram_font(23, bold=True))
    font = diagram_font(17)
    y = y1 + title_h + 12
    for attr in attributes:
        draw.text((x1 + 14, y), attr, font=font, fill="#111827")
        y += 23
    y = y1 + title_h + attr_h + 12
    for method in methods:
        draw.text((x1 + 14, y), method, font=font, fill="#374151")
        y += 23


def draw_lifeline(draw: ImageDraw.ImageDraw, x: int, title: str, y_top: int, y_bottom: int) -> None:
    font = diagram_font(22, bold=True)
    box = (x - 90, y_top, x + 90, y_top + 54)
    draw.rounded_rectangle(box, radius=8, outline=DIAGRAM_LINE, width=3, fill=DIAGRAM_FILL)
    draw_centered_text(draw, box, title, font)
    y = y_top + 54
    dash = 12
    while y < y_bottom:
        draw.line((x, y, x, min(y + dash, y_bottom)), fill="#94A3B8", width=3)
        y += dash * 2


def sequence_arrow(
    draw: ImageDraw.ImageDraw,
    x1: int,
    x2: int,
    y: int,
    label: str,
    reverse: bool = False,
) -> None:
    if x1 == x2:
        draw.line((x1, y, x1 + 90, y, x1 + 90, y + 45, x1 + 8, y + 45), fill=DIAGRAM_LINE, width=3)
        draw_arrow(draw, (x1 + 8, y + 45), (x1, y + 45), width=3)
        font = diagram_font(20)
        draw.text((x1 + 98, y + 12), label, font=font, fill="#111827")
        return
    start, end = ((x2, y), (x1, y)) if reverse else ((x1, y), (x2, y))
    draw_arrow(draw, start, end, width=3)
    font = diagram_font(20)
    w, h = text_size(draw, label, font)
    mx = (x1 + x2) // 2
    draw.rectangle((mx - w // 2 - 5, y - h - 10, mx + w // 2 + 5, y - 2), fill="white")
    draw.text((mx - w // 2, y - h - 8), label, font=font, fill="#111827")


def create_use_case_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1120), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "UML - diagram przypadkow uzycia", font=diagram_font(38, bold=True), fill="#111827")
    boundary = (360, 115, 1425, 1025)
    draw.rounded_rectangle(boundary, radius=28, outline="#111827", width=4, fill="#FAFBFC")
    draw.text((390, 135), "System: Budgetly", font=diagram_font(28, bold=True), fill="#111827")

    draw_actor(draw, 165, 175, "Gosc")
    draw_actor(draw, 165, 565, "Uzytkownik\nzalogowany")
    draw_actor(draw, 1600, 390, "Supabase\nAuth")

    use_cases = {
        "register": (455, 230, 790, 335, "Rejestracja /\nlogowanie"),
        "categories": (455, 420, 790, 525, "Zarzadzanie\nkategoriami"),
        "transactions": (455, 585, 790, 690, "CRUD\ntransakcji"),
        "dashboard": (850, 230, 1185, 335, "Panel\nmiesieczny"),
        "budgets": (850, 420, 1185, 525, "Budzety\ni alerty"),
        "recurring": (850, 585, 1185, 690, "Wydatki\ncykliczne"),
        "compare": (1040, 760, 1375, 865, "Porownanie\nmiesiecy"),
        "export": (650, 760, 985, 865, "Eksport CSV"),
        "settings": (455, 900, 790, 1000, "Ustawienia\nwaluty"),
    }
    for box in use_cases.values():
        draw_use_case(draw, box[:4], box[4])

    for key in ["register"]:
        x1, y1, x2, y2, _ = use_cases[key]
        draw_connector(draw, (210, 260), (x1, (y1 + y2) // 2))
        draw_connector(draw, (x2, (y1 + y2) // 2), (1510, 475), "<<include>>")
    for key in ["categories", "transactions", "dashboard", "budgets", "recurring", "compare", "export", "settings"]:
        x1, y1, x2, y2, _ = use_cases[key]
        draw_connector(draw, (220, 650), (x1, (y1 + y2) // 2))
    img.save(path)


def create_class_diagram(path: Path) -> None:
    img = Image.new("RGB", (2000, 1380), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "UML - diagram klas / model domeny", font=diagram_font(38, bold=True), fill="#111827")
    boxes = {
        "User": (90, 135, 450, 365),
        "Profile": (595, 135, 955, 365),
        "Category": (90, 560, 450, 820),
        "Transaction": (595, 520, 1015, 840),
        "Budget": (1160, 520, 1565, 840),
        "RecurringExpense": (595, 970, 1065, 1290),
    }

    relation_lines = [
        [(450, 250), (595, 250)],
        [(270, 365), (270, 560)],
        [(450, 320), (520, 320), (520, 520), (595, 590)],
        [(450, 285), (1360, 285), (1360, 520)],
        [(270, 365), (270, 455), (830, 455), (830, 970)],
        [(450, 650), (595, 650)],
        [(450, 740), (1160, 740)],
        [(450, 810), (520, 915), (595, 1070)],
    ]
    for points in relation_lines:
        draw.line(points, fill=DIAGRAM_LINE, width=3)

    draw_class_box(draw, boxes["User"], "User", ["+ id: uuid", "+ email: text", "+ created_at"], ["+ getSession()"])
    draw_class_box(draw, boxes["Profile"], "Profile", ["+ id: uuid", "+ user_id: uuid", "+ display_name", "+ currency_code"], ["+ updateCurrency()"])
    draw_class_box(draw, boxes["Category"], "Category", ["+ id: uuid", "+ user_id: uuid", "+ name", "+ color", "+ icon"], ["+ create()", "+ update()", "+ delete()"])
    draw_class_box(draw, boxes["Transaction"], "Transaction", ["+ id: uuid", "+ user_id: uuid", "+ category_id", "+ amount", "+ type", "+ transaction_date"], ["+ create()", "+ update()", "+ delete()"])
    draw_class_box(draw, boxes["Budget"], "Budget", ["+ id: uuid", "+ user_id: uuid", "+ category_id?", "+ limit_amount", "+ month", "+ year"], ["+ getProgress()", "+ getStatus()"])
    draw_class_box(draw, boxes["RecurringExpense"], "RecurringExpense", ["+ id: uuid", "+ user_id: uuid", "+ category_id", "+ frequency", "+ next_occurrence", "+ is_active"], ["+ processDue()", "+ toggleActive()"])

    def relation_label(x: int, y: int, text: str) -> None:
        font = diagram_font(22)
        w, h = text_size(draw, text, font)
        draw.rectangle((x - w // 2 - 6, y - h // 2 - 4, x + w // 2 + 6, y + h // 2 + 4), fill="white")
        draw.text((x - w // 2, y - h // 2), text, font=font, fill="#111827")

    relation_label(520, 250, "1  do  1")
    relation_label(270, 455, "1  do  *")
    relation_label(545, 435, "1  do  *")
    relation_label(980, 285, "1  do  *")
    relation_label(805, 455, "1  do  *")
    relation_label(520, 650, "1  do  *")
    relation_label(805, 740, "0..1  do  *")
    relation_label(535, 930, "1  do  *")
    img.save(path)


def create_sequence_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 1280), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "UML - diagram sekwencji: dodanie transakcji", font=diagram_font(38, bold=True), fill="#111827")
    xs = [160, 465, 775, 1085, 1395, 1690]
    titles = ["Uzytkownik", "Formularz UI", "Server Action", "Service", "Supabase", "PostgreSQL/RLS"]
    for x, title in zip(xs, titles):
        draw_lifeline(draw, x, title, 125, 1185)

    y = 235
    sequence_arrow(draw, xs[0], xs[1], y, "wypelnia i wysyla formularz")
    y += 82
    sequence_arrow(draw, xs[1], xs[2], y, "createTransactionAction(formData)")
    y += 82
    sequence_arrow(draw, xs[2], xs[2], y, "walidacja Zod")
    y += 82
    sequence_arrow(draw, xs[2], xs[4], y, "auth.getUser()")
    y += 82
    sequence_arrow(draw, xs[2], xs[3], y, "createTransaction(userId, input)")
    y += 82
    sequence_arrow(draw, xs[3], xs[4], y, "insert transactions")
    y += 82
    sequence_arrow(draw, xs[4], xs[5], y, "INSERT + policy RLS")
    y += 82
    sequence_arrow(draw, xs[5], xs[4], y, "rekord albo blad")
    y += 82
    sequence_arrow(draw, xs[4], xs[3], y, "data/error")
    y += 82
    sequence_arrow(draw, xs[3], xs[2], y, "Transaction")
    y += 82
    sequence_arrow(draw, xs[2], xs[1], y, "revalidatePath + wynik")
    img.save(path)


def create_component_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 1180), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "UML - diagram komponentow", font=diagram_font(38, bold=True), fill="#111827")
    components = {
        "browser": (80, 180, 390, 330, "Browser UI", "React components\nforms, charts, nav"),
        "router": (545, 160, 900, 340, "Next.js App Router", "pages, layouts,\nserver rendering"),
        "actions": (1040, 160, 1395, 340, "Server Actions", "controllers,\nvalidation, redirects"),
        "services": (1040, 475, 1395, 655, "Services", "business logic,\nqueries, aggregation"),
        "supabase": (1515, 430, 1835, 610, "Supabase Client", "SSR/browser client"),
        "auth": (1515, 145, 1835, 305, "Supabase Auth", "sessions,\nusers"),
        "db": (1515, 760, 1835, 945, "PostgreSQL + RLS", "tables, indexes,\npolicies"),
        "fx": (545, 760, 900, 930, "Frankfurter API", "exchange rates\ncache fallback"),
        "vercel": (80, 760, 390, 930, "Vercel", "build and deploy"),
    }
    for box in components.values():
        draw_component(draw, box[:4], box[4], box[5])

    draw_arrow(draw, (390, 255), (545, 255))
    draw_arrow(draw, (900, 250), (1040, 250))
    draw_arrow(draw, (1218, 340), (1218, 475))
    draw_arrow(draw, (1395, 565), (1515, 525))
    draw_arrow(draw, (1675, 610), (1675, 760))
    draw_arrow(draw, (1395, 250), (1515, 225))
    draw_arrow(draw, (722, 760), (1040, 625))
    draw_arrow(draw, (235, 760), (235, 330))
    img.save(path)


def build_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(exist_ok=True)
    paths = {
        "use_case": DIAGRAM_DIR / "uml_use_case_budgetly.png",
        "class": DIAGRAM_DIR / "uml_class_budgetly.png",
        "sequence": DIAGRAM_DIR / "uml_sequence_add_transaction.png",
        "component": DIAGRAM_DIR / "uml_component_budgetly.png",
    }
    create_use_case_diagram(paths["use_case"])
    create_class_diagram(paths["class"])
    create_sequence_diagram(paths["sequence"])
    create_component_diagram(paths["component"])
    return paths


def add_diagram_image(doc: Document, image_path: Path, caption: str, alt_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(6.3))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", polish_text(alt_text))
    doc_pr.set("title", polish_text(caption))

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    add_text(caption_p, caption, italic=True, color=GRAY, size=9.5)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = width_to_inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_row_as_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def width_to_inches(dxa: int):
    return Inches(dxa / 1440)


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_text(paragraph, text: str, bold=False, italic=False, color=None, size=None):
    text = polish_text(text)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    return run


def add_para(doc: Document, text: str = "", style: str | None = None, after=6, before=0):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if text:
        add_text(paragraph, text)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    text = polish_text(text)
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = add_para(doc, style="List Bullet", after=4)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    add_text(paragraph, text)
    return paragraph


def add_numbered(doc: Document, text: str):
    paragraph = add_para(doc, style="List Number", after=4)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    add_text(paragraph, text)
    return paragraph


def fill_cell(cell, text: str, bold=False, color=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    if align is not None:
        paragraph.alignment = align
    add_text(paragraph, text, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_row_as_header(table.rows[0])
    for index, header in enumerate(headers):
        fill_cell(table.rows[0].cells[index], header, bold=True, color=BLACK)
        set_cell_shading(table.rows[0].cells[index], header_fill)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            fill_cell(row.cells[index], value)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent=120)
    mark_row_as_header(table.rows[0])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, f"{title}: ", bold=True, color=DARK_BLUE)
    add_text(p, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer_p, "Budgetly - dokumentacja projektowa", color=GRAY, size=9)


def add_front_matter(doc: Document) -> None:
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "Programowanie aplikacji internetowych", color=GRAY, size=9)

    kicker = add_para(doc, "DOKUMENTACJA PROJEKTOWA", after=4)
    add_text(kicker, "", bold=True)
    kicker.runs[0].font.color.rgb = GRAY
    kicker.runs[0].font.size = Pt(10)

    title = add_para(doc, after=4)
    add_text(title, "Budgetly", bold=True, color=BLACK, size=28)

    subtitle = add_para(doc, "Aplikacja webowa do zarzadzania budzetem osobistym", after=16)
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = GRAY

    meta_rows = [
        ("Przedmiot", "Programowanie aplikacji internetowych"),
        ("Projekt", "Budgetly - wielouzytkownikowy dziennik wydatkow"),
        (
            "Autorzy",
            "Karol Scislowski, 177306, Informatyka Online, K38; Jakub Grygiel, 174556, Informatyka Online, K38",
        ),
        ("Data dokumentu", "30.05.2026"),
        ("Wersja", "1.0"),
        ("Zakres", "Dokumentacja wstepna, specyfikacja, harmonogram, architektura, implementacja, testowanie i plan demonstracji"),
    ]
    add_table(doc, ["Metadane", "Wartosc"], meta_rows, [2100, 7260], header_fill=BLUE_GRAY)

    rule = add_para(doc, after=10)
    set_paragraph_border_bottom(rule)

    add_callout(
        doc,
        "Cel dokumentu",
        "Dokument zostal przygotowany tak, aby pokrywal kryteria oceniania z pliku 'Programowanie aplikacji internetowych - plan przedmiotu oraz zasady zaliczenia przedmiotu.pdf'. Kolejnosc rozdzialow odpowiada glownym grupom punktowym z kryteriow.",
    )


def section_intro(doc: Document) -> None:
    add_heading(doc, "1. Dokumentacja wstepna", 1)
    add_heading(doc, "1.1 Koncepcja projektu", 2)
    add_para(
        doc,
        "Budgetly to wielouzytkownikowa aplikacja webowa do ewidencji finansow osobistych. System pozwala rejestrowac przychody i wydatki, kategoryzowac transakcje, ustawiac budzety miesieczne, automatyzowac wydatki cykliczne, porownywac miesiace oraz eksportowac dane do pliku CSV. Dane kazdego uzytkownika sa odseparowane kontem Supabase Auth i politykami Row Level Security w bazie PostgreSQL.",
    )

    add_heading(doc, "1.2 Mierzalne cele systemu", 2)
    add_table(
        doc,
        ["Cel", "Miernik sukcesu", "Uzasadnienie"],
        [
            [
                "Szybka rejestracja operacji finansowej",
                "Dodanie poprawnej transakcji przez zalogowanego uzytkownika w czasie do 30 sekund.",
                "Najczestsza czynnosc musi byc szybka, aby uzytkownik regularnie uzupelnial dane.",
            ],
            [
                "Biezacy wglad w sytuacje finansowa",
                "Uzyskanie salda, sumy przychodow i wydatkow wybranego miesiaca na panelu w czasie do 2 minut od zalogowania.",
                "System ma odpowiadac na podstawowe pytanie: ile wydano, ile zarobiono i jaki jest bilans.",
            ],
            [
                "Kontrola przekroczen budzetu",
                "Dla kazdego budzetu miesiecznego system pokazuje status ok, warning lub exceeded oraz procent wykorzystania limitu.",
                "Uzytkownik otrzymuje sygnal ostrzegawczy zanim utraci kontrole nad limitem wydatkow.",
            ],
        ],
        [2300, 3300, 3760],
        header_fill=BLUE_GRAY,
    )

    add_heading(doc, "1.3 Interesariusze", 2)
    add_table(
        doc,
        ["Interesariusz", "Potrzeba / interes"],
        [
            ["Uzytkownik koncowy", "Chce kontrolowac wydatki, budzety i trendy finansowe bez arkuszy kalkulacyjnych."],
            ["Prowadzacy / oceniający projekt", "Ocenia kompletność analizy, spojnosc architektury, jakosc implementacji i demonstracje."],
            ["Autor / zespol projektowy", "Potrzebuje czytelnego zakresu, harmonogramu, kryteriow akceptacji i planu testow."],
            ["Administrator techniczny / utrzymaniowiec", "Odpowiada za konfiguracje Supabase, Vercel, zmienne srodowiskowe i diagnostyke bledow."],
        ],
        [3000, 6360],
    )

    add_heading(doc, "1.4 Szczegolna wartosc dla uzytkownikow koncowych", 2)
    for item in [
        "Centralne miejsce dla transakcji, kategorii, budzetow i porownan miesiecy zamiast rozproszonych notatek lub arkuszy.",
        "Automatyczne generowanie wydatkow cyklicznych ogranicza pomijanie stalych kosztow, np. abonamentow i rachunkow.",
        "Alerty budzetowe i wykresy kategorii pomagaja szybko rozpoznac obszary nadmiernych wydatkow.",
        "Eksport CSV pozwala zachowac kopie danych lub wykonac dodatkowa analize poza aplikacja.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.5 Docelowa platforma technologiczna", 2)
    add_table(
        doc,
        ["Platforma", "Uzasadnienie wyboru"],
        [
            [
                "Responsywna aplikacja webowa",
                "Uzytkownik moze korzystac z systemu na komputerze i telefonie bez instalowania aplikacji natywnej. Next.js App Router obsluguje widoki publiczne, autoryzowane dashboardy i serwerowe pobieranie danych.",
            ],
            [
                "Backend cloud: Supabase",
                "Supabase dostarcza Auth, PostgreSQL i RLS, co skraca czas implementacji bez rezygnacji z relacyjnego modelu danych i kontroli dostepu.",
            ],
            [
                "Deploy cloud: Vercel",
                "Vercel dobrze pasuje do aplikacji Next.js, automatyzuje build i pozwala latwo udostepnic wersje demonstracyjna projektu.",
            ],
        ],
        [2600, 6760],
    )

    add_heading(doc, "1.6 Zagrozenia i ograniczenia projektowe", 2)
    add_table(
        doc,
        ["Ryzyko / ograniczenie", "Wplyw", "Ograniczenie ryzyka"],
        [
            [
                "Bledna konfiguracja Supabase lub RLS",
                "Ryzyko braku dostepu do danych albo naruszenia izolacji danych uzytkownikow.",
                "Migracje SQL wlaczaja RLS dla tabel publicznych; test demo powinien uzyc dwoch kont i sprawdzic separacje danych.",
            ],
            [
                "Zaleznosc od zewnetrznego API kursow walut",
                "Brak aktualnego kursu moze znieksztalcic wyswietlane wartosci w walucie innej niz PLN.",
                "Kod stosuje cache i fallback do ostatniej znanej wartosci lub kursu 1.",
            ],
            [
                "Jakosc danych wprowadzanych recznie",
                "Niepelne lub zle kategorie obnizaja wartosc analiz i porownan.",
                "Formularze sa walidowane schematami Zod, a domyslne kategorie sa seedowane po rejestracji.",
            ],
            [
                "Ograniczony czas projektu zaliczeniowego",
                "Nie wszystkie funkcje aplikacji finansowej, np. OCR paragonow, sa realistyczne w zakresie projektu.",
                "Zakres skupia sie na podstawowym cyklu: zapis danych, analiza, budzety, eksport i prezentacja.",
            ],
        ],
        [2600, 3000, 3760],
    )


def section_spec(doc: Document, diagrams: dict[str, Path]) -> None:
    add_heading(doc, "2. Zarys specyfikacji", 1)
    add_heading(doc, "2.1 Aktorzy systemu", 2)
    add_table(
        doc,
        ["Aktor", "Rola w systemie"],
        [
            ["Gosc niezalogowany", "Oglada strone startowa, przechodzi do logowania lub rejestracji."],
            ["Uzytkownik zalogowany", "Zarzadza wlasnymi kategoriami, transakcjami, budzetami, wydatkami cyklicznymi, waluta i eksportem."],
            ["Supabase Auth", "Zewnetrzny komponent uwierzytelnienia; zaklada sesje i identyfikuje uzytkownika dla RLS."],
            ["Procesor wydatkow cyklicznych", "Wewnetrzny aktor systemowy uruchamiany w layoucie dashboardu, generujacy zalegle transakcje cykliczne."],
            ["Administrator techniczny", "Konfiguruje zmienne srodowiskowe, migracje, deploy i podstawowa diagnostyke."],
        ],
        [3000, 6360],
    )
    add_heading(doc, "2.1.1 Diagram UML przypadkow uzycia", 3)
    add_para(
        doc,
        "Diagram pokazuje podstawowe interakcje aktorow z systemem Budgetly: obsluge konta, zarzadzanie finansami, analityke, eksport oraz zaleznosc od Supabase Auth.",
    )
    add_diagram_image(
        doc,
        diagrams["use_case"],
        "Rysunek 1. Diagram UML przypadkow uzycia systemu Budgetly.",
        "Diagram przypadkow uzycia: gosc korzysta z rejestracji i logowania, uzytkownik zalogowany zarzadza kategoriami, transakcjami, budzetami, wydatkami cyklicznymi, dashboardem, porownaniem miesiecy, eksportem CSV i ustawieniami waluty.",
    )

    add_heading(doc, "2.2 Obiekty dziedzinowe", 2)
    add_table(
        doc,
        ["Obiekt", "Najwazniejsze atrybuty", "Znaczenie"],
        [
            ["Profile", "user_id, display_name, currency_code", "Ustawienia uzytkownika i waluta prezentacji."],
            ["Category", "name, color, icon, user_id", "Klasyfikacja transakcji i budzetow."],
            ["Transaction", "amount, type, date, category_id, description", "Podstawowy zapis przychodu lub wydatku."],
            ["Budget", "limit_amount, month, year, category_id", "Limit wydatkow ogolny lub dla kategorii."],
            ["RecurringExpense", "amount, frequency, next_occurrence, is_active", "Definicja powtarzalnego kosztu generowanego automatycznie."],
            ["DashboardData", "summary, categoryBreakdown, recentTransactions", "Zagregowany widok miesiaca dla panelu."],
            ["MonthComparisonResult", "period1, period2, categories, difference", "Porownanie wydatkow miedzy dwoma miesiacami."],
            ["CsvTransactionRow", "transaction_date, description, amount, type, category_name", "Rekord eksportu danych do CSV."],
        ],
        [1900, 3400, 4060],
    )
    add_heading(doc, "2.2.1 Diagram UML klas domenowych", 3)
    add_para(
        doc,
        "Diagram klas opisuje najwazniejsze encje dziedzinowe oraz relacje wynikajace bezposrednio ze schematu Supabase i serwisow aplikacji.",
    )
    add_diagram_image(
        doc,
        diagrams["class"],
        "Rysunek 2. Diagram UML klas / model domeny Budgetly.",
        "Diagram klas: User ma profil, kategorie, transakcje, budzety i wydatki cykliczne; Category jest powiazana z Transaction, Budget i RecurringExpense.",
    )

    add_heading(doc, "2.3 Kontekst systemu i glowne powiazania", 2)
    add_callout(
        doc,
        "Diagram tekstowy",
        "Gosc -> Rejestracja/Logowanie -> Supabase Auth -> Sesja; Uzytkownik -> Next.js App Router -> Server Actions -> Services -> Supabase Client -> PostgreSQL/RLS; Dashboard -> Services -> Agregacje miesieczne; Procesor cykliczny -> RecurringExpense -> Transaction; Uzytkownik -> Eksport CSV -> CsvTransactionRow.",
    )
    add_table(
        doc,
        ["Powiazanie / czynnosc", "Aktor", "Obiekt / modul"],
        [
            ["Rejestracja konta i utworzenie profilu", "Gosc niezalogowany", "Supabase Auth, Profile, seedDefaultCategories"],
            ["Dodanie transakcji", "Uzytkownik zalogowany", "Transaction, Category, transactionSchema"],
            ["Obliczenie podsumowania miesiaca", "Uzytkownik zalogowany", "DashboardData, Transaction"],
            ["Kontrola budzetu", "Uzytkownik zalogowany", "Budget, Transaction, getBudgetProgress"],
            ["Wygenerowanie zaleglych kosztow", "Procesor wydatkow cyklicznych", "RecurringExpense, Transaction"],
            ["Eksport danych", "Uzytkownik zalogowany", "CsvTransactionRow, generateTransactionsCsv"],
        ],
        [3400, 2500, 3460],
    )

    add_heading(doc, "2.4 Wymagania funkcjonalne", 2)
    functional_rows = [
        ["F1", "Rejestracja, logowanie i wylogowanie", "Uzytkownik moze utworzyc konto, zalogowac sie i zakonczyc sesje.", "Poprawny email i haslo tworza konto lub sesje; trasy chronione przekierowuja niezalogowanego uzytkownika do /login."],
        ["F2", "CRUD kategorii", "Uzytkownik moze dodac, edytowac i usunac wlasne kategorie z kolorem oraz ikona.", "Nazwa, kolor hex i ikona sa walidowane; lista kategorii pokazuje tylko rekordy user_id aktualnego uzytkownika."],
        ["F3", "CRUD transakcji i filtrowanie", "Uzytkownik zarzadza wydatkami i przychodami oraz filtruje je po dacie, kategorii i typie.", "Kwota musi byc dodatnia, data poprawna, kategoria istnieje; filtry modyfikuja liste bez pokazywania cudzych danych."],
        ["F4", "Panel miesieczny", "System pokazuje przychody, wydatki, saldo, liczbe transakcji, podzial kategorii i ostatnie transakcje.", "Dane sa liczone dla wybranego miesiaca i roku, a brak danych daje wartosci zerowe i puste listy."],
        ["F5", "Budzetowanie", "Uzytkownik definiuje limit miesieczny ogolny lub dla kategorii.", "System oblicza procent wykorzystania i status ok/warning/exceeded; duplikat budzetu dla tej samej kategorii i miesiaca jest blokowany."],
        ["F6", "Wydatki cykliczne", "Uzytkownik definiuje powtarzalne koszty tygodniowe, miesieczne lub roczne.", "Aktywne rekordy z next_occurrence <= dzis generuja transakcje i przesuwaja nastepna date wystapienia."],
        ["F7", "Porownanie miesiecy", "Uzytkownik wybiera dwa okresy i porownuje wydatki wedlug kategorii.", "Wynik zawiera sumy okresow, roznice per kategoria oraz wykres slupkowy i tabele roznic."],
        ["F8", "Eksport CSV", "Uzytkownik pobiera transakcje do pliku CSV.", "Plik zawiera naglowek Data, Opis, Kwota, Typ, Kategoria; wartosci z przecinkami i cudzyslowami sa escapowane."],
        ["F9", "Waluta prezentacji", "Uzytkownik wybiera walute wyswietlania kwot.", "Dozwolone sa tylko wspierane kody walut; wartosci sa konwertowane z PLN przez cache kursow."],
        ["F10", "Tryb ciemny i responsywny layout", "Interfejs dziala na telefonie i desktopie oraz wspiera dark mode.", "Nawigacja dolna jest dostepna na mobile, sidebar na szerszych ekranach, a ustawienie motywu zmienia wyglad ekranow."],
    ]
    add_table(doc, ["ID", "Wymaganie", "Opis", "Kryteria akceptacji"], functional_rows, [600, 2100, 3000, 3660], header_fill=BLUE_GRAY)

    add_heading(doc, "2.5 Wymagania niefunkcjonalne", 2)
    add_table(
        doc,
        ["ID", "Wymaganie", "Kryteria akceptacji"],
        [
            ["NF1", "Bezpieczenstwo i izolacja danych", "Kazda tabela publiczna ma wlaczone RLS; zapytania serwisowe filtrują po user_id; test demonstracyjny potwierdza brak dostepu do danych innego konta."],
            ["NF2", "Poprawnosc walidacji", "Formularze odrzucaja niepoprawne UUID kategorii, kwoty <= 0, bledne daty, niepoprawne kolory hex i nieobslugiwane waluty."],
            ["NF3", "Uzytecznosc i responsywnosc", "Glowne ekrany sa czytelne na szerokosciach ok. 375 px, 768 px i 1280 px; podstawowy przeplyw dodania transakcji nie wymaga przejscia przez wiele ekranow."],
            ["NF4", "Utrzymywalnosc", "Logika biznesowa jest w services, walidacja w lib/validations, a akcje serwerowe pelnia role kontrolerow; nowe moduly mozna testowac jednostkowo bez UI."],
        ],
        [700, 2500, 6160],
    )


def section_schedule_team(doc: Document) -> None:
    add_heading(doc, "3. Harmonogram i zespol - zarys", 1)
    add_heading(doc, "3.1 Ramowy harmonogram realizacji", 2)
    add_table(
        doc,
        ["Etap", "Zakres prac", "Rezultat"],
        [
            ["1. Analiza potrzeb", "Okreslenie koncepcji, aktorow, obiektow, ryzyk i wymagan.", "Zarys specyfikacji i zakres MVP."],
            ["2. Projekt danych", "Model relacyjny, migracje SQL, RLS, profile i waluta.", "Schemat PostgreSQL w Supabase."],
            ["3. Uwierzytelnienie i layout", "Rejestracja, logowanie, middleware, dashboard layout, nawigacja.", "Chronione trasy i sesja uzytkownika."],
            ["4. Moduly CRUD", "Kategorie, transakcje, budzety, wydatki cykliczne.", "Podstawowa funkcjonalnosc biznesowa."],
            ["5. Analityka i eksport", "Dashboard, porownanie miesiecy, wykresy, CSV, waluty.", "Funkcje wspierajace decyzje uzytkownika."],
            ["6. Testy i stabilizacja", "Testy walidacji, funkcji serwisowych, eksportu, walut; poprawki bledow.", "Przechodzacy zestaw testow i build."],
            ["7. Dokumentacja i demonstracja", "README, diagram bazy, checklista demo, niniejsza dokumentacja.", "Material do oddania i scenariusz prezentacji."],
        ],
        [1600, 4700, 3060],
        header_fill=BLUE_GRAY,
    )

    add_heading(doc, "3.2 Przykladowy sklad zespolu projektowego", 2)
    add_para(
        doc,
        "Projekt moze byc realizowany indywidualnie, natomiast ponizszy sklad jest hipotetycznym podzialem rol wymaganym do planowania prac projektowych.",
    )
    add_table(
        doc,
        ["Rola", "Zakres odpowiedzialnosci"],
        [
            ["Analityk systemowy / Product Owner", "Definicja problemu, celow, wymagan, kryteriow akceptacji i priorytetow demonstracji."],
            ["Projektant UX/UI", "Projekt przeplywow ekranow, responsywna nawigacja, czytelnosc formularzy, dark mode."],
            ["Frontend Developer", "Implementacja widokow Next.js, komponentow React, formularzy i integracji z akcjami serwerowymi."],
            ["Backend / Database Developer", "Migracje Supabase, RLS, serwisy domenowe, relacje i indeksy bazy danych."],
            ["QA / Tester", "Scenariusze testowe, testy jednostkowe, walidacja formularzy, checklista demo i testy regresji."],
            ["DevOps / Dokumentalista", "Konfiguracja Vercel, zmiennych srodowiskowych, README, dokumentacja projektowa i instrukcje uruchomienia."],
        ],
        [3000, 6360],
    )


def section_architecture(doc: Document, diagrams: dict[str, Path]) -> None:
    add_heading(doc, "4. Planowanie architektury, implementacji i testowania", 1)
    add_heading(doc, "4.1 Stos technologiczny", 2)
    add_table(
        doc,
        ["Element stosu", "Zastosowanie", "Uzasadnienie"],
        [
            ["Next.js 16 App Router", "Routing, layouty, strony publiczne i chronione.", "Pozwala laczyc renderowanie serwerowe, Server Actions i strukture modulow aplikacji webowej."],
            ["React 19 + TypeScript", "Komponenty UI i typowanie danych.", "Zmniejsza ryzyko bledow kontraktow miedzy formularzami, serwisami i widokami."],
            ["Supabase Auth + PostgreSQL + RLS", "Uwierzytelnienie, baza danych i izolacja rekordow.", "Dostarcza gotowa autoryzacje oraz relacyjny model danych z politykami bezpieczenstwa."],
            ["Tailwind CSS 4 + shadcn/ui", "Warstwa prezentacji i spójne komponenty.", "Przyspiesza budowe responsywnego interfejsu w jezyku polskim."],
            ["Zod", "Walidacja formularzy i filtrow.", "Jednoznaczne kryteria poprawnosci danych wejsciowych przed zapisem do bazy."],
            ["Recharts", "Wykresy kategorii i porownania miesiecy.", "Czytelna wizualizacja danych finansowych bez wlasnej implementacji chartingu."],
            ["Vitest", "Automatyczne testy jednostkowe.", "Szybka weryfikacja logiki serwisowej, walidacji, eksportu i formatowania walut."],
            ["Vercel", "Build i publikacja aplikacji.", "Naturalne srodowisko deploy dla Next.js z obsluga zmiennych srodowiskowych."],
        ],
        [2200, 3300, 3860],
        header_fill=BLUE_GRAY,
    )

    add_heading(doc, "4.2 Ogolna architektura aplikacji", 2)
    add_para(
        doc,
        "Budgetly ma architekture modulowego monolitu webowego. Frontend i warstwa akcji serwerowych znajduja sie w jednej aplikacji Next.js, natomiast uwierzytelnienie, baza danych i polityki dostepu sa realizowane przez Supabase. Taki wybor ogranicza narzut infrastrukturalny projektu zaliczeniowego, ale zachowuje jasny podzial odpowiedzialnosci.",
    )
    add_callout(
        doc,
        "Schemat warstw",
        "Browser UI -> Next.js routes/layouts -> Server Actions -> Services -> Supabase client -> PostgreSQL + RLS. Middleware chroni trasy, a komponenty UI korzystaja z danych przygotowanych przez strony serwerowe i akcje.",
    )
    add_diagram_image(
        doc,
        diagrams["component"],
        "Rysunek 3. Diagram UML komponentow aplikacji Budgetly.",
        "Diagram komponentow: Browser UI komunikuje sie z Next.js App Router, Server Actions, Services, Supabase Client, Supabase Auth, PostgreSQL z RLS, API kursow walut i Vercel.",
    )
    add_table(
        doc,
        ["Komponent / modul", "Odpowiedzialnosc"],
        [
            ["app/", "Trasy Next.js, layouty publiczne i dashboardowe, pobieranie danych dla stron."],
            ["components/", "Komponenty prezentacyjne i formularze: transakcje, budzety, kategorie, dashboard, ustawienia."],
            ["actions/", "Server Actions jako kontrolery: autoryzacja, walidacja wejscia, wywolanie serwisow i revalidatePath."],
            ["services/", "Logika biznesowa i zapytania: transakcje, budzety, cykliczne, porownania, eksport, waluty."],
            ["lib/validations/", "Schematy Zod dla formularzy i filtrow."],
            ["lib/supabase/", "Klient serwerowy, przegladarkowy, middleware i admin client dla dev."],
            ["supabase/migrations/", "Definicje tabel, enumow, indeksow, trigger profilu i polityk RLS."],
            ["tests (*.test.ts)", "Testy jednostkowe logiki, walidacji, eksportu i formatowania pieniedzy."],
        ],
        [2600, 6760],
    )

    add_heading(doc, "4.3 Wzorce projektowe", 2)
    add_table(
        doc,
        ["Wzorzec", "Przyklad w projekcie", "Uzasadnienie"],
        [
            [
                "Layered Architecture / Service Layer",
                "app -> actions -> services -> Supabase; np. transaction.actions.ts wywoluje transaction.service.ts.",
                "Oddziela UI, kontrolery, logike biznesowa i persystencje. Ulatwia testowanie funkcji takich jak getBudgetStatus, buildCategoryComparison i generateTransactionsCsv.",
            ],
            [
                "Provider / Context",
                "CurrencyProvider przekazuje walute i kurs do komponentow dashboardu.",
                "Zmniejsza duplikacje przekazywania ustawien waluty przez wiele poziomow komponentow i pozwala utrzymac spójny format kwot.",
            ],
            [
                "Middleware Guard",
                "middleware.ts i lib/supabase/middleware.ts przekierowuja niezalogowanych z tras chronionych.",
                "Centralizuje ochrone dostepu zamiast powielac sprawdzanie sesji w kazdym ekranie.",
            ],
        ],
        [2300, 3300, 3760],
        header_fill=BLUE_GRAY,
    )

    add_heading(doc, "4.4 Fragment warstwy implementacyjnej: baza danych", 2)
    add_table(
        doc,
        ["Tabela / element", "Kluczowe pola i ograniczenia", "Relacje / uwagi"],
        [
            ["auth.users", "id, email, created_at", "Tabela Supabase Auth; zrodlo tozsamosci uzytkownika."],
            ["profiles", "user_id UNIQUE, display_name, currency_code CHECK", "1:1 z auth.users; trigger handle_new_user tworzy profil."],
            ["categories", "user_id, name, color, icon", "1:N z users; FK do transactions, budgets i recurring_expenses."],
            ["transactions", "amount CHECK > 0, type enum, date, category_id", "Podstawowy zapis operacji; indeksy po user_id, dacie i kategorii."],
            ["budgets", "limit_amount CHECK > 0, month 1-12, year >= 2020", "Limit ogolny lub kategorii; UNIQUE(user_id, category_id, month, year)."],
            ["recurring_expenses", "frequency enum, next_occurrence, is_active", "Zrodlo automatycznie generowanych transakcji."],
            ["transaction_type", "expense, income", "Enum ograniczajacy typ transakcji."],
            ["frequency_type", "weekly, monthly, yearly", "Enum ograniczajacy czestotliwosc kosztow cyklicznych."],
            ["RLS policies", "SELECT/INSERT/UPDATE/DELETE z auth.uid() = user_id", "Izolacja danych miedzy kontami uzytkownikow."],
        ],
        [2100, 3900, 3360],
    )
    add_heading(doc, "4.4.1 Diagram UML sekwencji dla dodania transakcji", 3)
    add_para(
        doc,
        "Sekwencja pokazuje reprezentatywny przeplyw implementacyjny: formularz po stronie UI przekazuje dane do Server Action, walidacja Zod poprzedza wywolanie serwisu, a zapis konczy sie operacja INSERT w PostgreSQL z kontrola RLS.",
    )
    add_diagram_image(
        doc,
        diagrams["sequence"],
        "Rysunek 4. Diagram UML sekwencji: dodanie transakcji.",
        "Diagram sekwencji dodania transakcji: uzytkownik wysyla formularz, Server Action waliduje dane, pobiera sesje, wywoluje service, Supabase wykonuje insert do PostgreSQL z RLS i zwraca wynik do interfejsu.",
    )

    add_heading(doc, "4.5 Organizacja pracy zespolu developerskiego", 2)
    add_table(
        doc,
        ["Element organizacji", "Uzasadnienie"],
        [
            ["Podzial wedlug warstw i modulow", "Struktura app/actions/services/lib/supabase pozwala rownolegle pracowac nad UI, logika biznesowa i baza danych bez mieszania odpowiedzialnosci."],
            ["Migracje bazy jako artefakty projektu", "Zmiany schematu sa odtwarzalne i mozna je uruchomic kolejno w Supabase SQL Editor."],
            ["Definition of Done dla funkcji", "Funkcja jest gotowa, gdy ma walidacje danych, obsluge bledow, test lub scenariusz manualny oraz wpis w checkliscie demo."],
            ["Kontrola jakosci przed oddaniem", "Przed demonstracja uruchamiane sa npm test, npm run build oraz manualny przechod przez najwazniejsze ekrany."],
        ],
        [3200, 6160],
        header_fill=BLUE_GRAY,
    )

    add_heading(doc, "4.6 Plan testowania", 2)
    add_table(
        doc,
        ["Technika", "Sposob testowania", "Znaczenie w procesie"],
        [
            ["Testy jednostkowe logiki biznesowej", "Automatyczne, Vitest; m.in. getBudgetStatus, calculateNextOccurrence, buildCategoryComparison.", "Chronia reguly domenowe przed regresja podczas zmian w UI lub bazie."],
            ["Testy walidacji schematow", "Automatyczne, Vitest + Zod; transakcje, kategorie, auth, waluty.", "Zapewniaja, ze system odrzuca niepoprawne dane zanim trafią do bazy."],
            ["Testy eksportu i formatowania", "Automatyczne, Vitest; CSV escaping i formatowanie/konwersja walut.", "Minimalizuja ryzyko blednych plikow eksportu i kwot prezentowanych uzytkownikowi."],
            ["Testy integracyjne manualne z Supabase", "Manualne; rejestracja, logowanie, CRUD, RLS, migracje i seed danych demo.", "Sprawdzaja przeplyw przez realne uslugi chmurowe i konfiguracje srodowiska."],
            ["Testy responsywnosci i UX", "Manualne; szerokosci ok. 375 px, 768 px, 1280 px oraz dark mode.", "Weryfikuja, czy aplikacja jest uzywalna na telefonie i desktopie."],
            ["Build produkcyjny", "Automatyczne polecenie npm run build.", "Potwierdza zgodnosc kodu z kompilacja Next.js i gotowosc do deployu."],
        ],
        [2500, 3300, 3560],
    )


def section_demo_and_matrix(doc: Document) -> None:
    add_heading(doc, "5. Plan demonstracji", 1)
    add_para(
        doc,
        "Punktacja za demonstracje jest przyznawana podczas prezentacji, dlatego dokument zawiera proponowany scenariusz pokazu oraz miejsca, w ktorych nalezy omowic elementy analizy i projektu.",
    )
    demo_steps = [
        "Pokaz strony startowej i wyjasnienie celu Budgetly.",
        "Rejestracja lub logowanie oraz omowienie Supabase Auth i chronionych tras.",
        "Dodanie kategorii i transakcji, pokaz walidacji formularzy.",
        "Panel miesieczny: saldo, wydatki, przychody, wykres kategorii i ostatnie transakcje.",
        "Budzet miesieczny: ustawienie limitu, status warning/exceeded i alert.",
        "Wydatki cykliczne: definicja kosztu i omowienie automatycznego generowania.",
        "Porownanie dwoch miesiecy oraz eksport CSV.",
        "Krotkie omowienie architektury: app/actions/services/Supabase/RLS.",
        "Podsumowanie testow, README, diagramu bazy i checklisty demo.",
    ]
    for step in demo_steps:
        add_numbered(doc, step)

    add_heading(doc, "6. Macierz spelnienia kryteriow oceniania", 1)
    add_table(
        doc,
        ["Kryterium z PDF", "Miejsce w dokumencie", "Status"],
        [
            ["1a. Krotki opis koncepcji projektu", "1.1", "Uwzgledniono"],
            ["1b. Co najmniej dwa mierzalne cele", "1.2", "3 cele"],
            ["1c. Co najmniej dwoch interesariuszy", "1.3", "4 interesariuszy"],
            ["1d. Co najmniej dwie wartosci dla uzytkownikow", "1.4", "4 przyklady"],
            ["1e. Platforma technologiczna z uzasadnieniem", "1.5", "3 platformy/warstwy"],
            ["1f. Co najmniej dwa ryzyka/ograniczenia", "1.6", "4 ryzyka"],
            ["2a. Aktorzy systemu", "2.1", "5 aktorow"],
            ["2b. Obiekty systemu", "2.2", "8 obiektow"],
            ["2c. Czynnosci i powiazania", "2.1.1, 2.3", "Diagram UML przypadkow uzycia + 6 powiazan"],
            ["2d. Co najmniej 7 wymagan funkcjonalnych z akceptacja", "2.4", "10 wymagan"],
            ["2e. Co najmniej 2 wymagania niefunkcjonalne z akceptacja", "2.5", "4 wymagania"],
            ["3a. Harmonogram min. 5 wpisow", "3.1", "7 etapow"],
            ["3b. Sklad zespolu i odpowiedzialnosci", "3.2", "6 rol"],
            ["4a. Stos technologiczny z uzasadnieniem", "4.1", "8 elementow"],
            ["4b. Architektura i moduly", "4.2", "Diagram UML komponentow + 8 modulow"],
            ["4c. Wzorzec projektowy z uzasadnieniem", "4.3", "3 wzorce"],
            ["4d. Warstwa implementacyjna / baza danych", "2.2.1, 4.4, 4.4.1", "Diagram klas, diagram sekwencji i 9 elementow"],
            ["4e. Organizacja pracy zespolu developerskiego", "4.5", "4 elementy"],
            ["4f. Co najmniej 3 techniki testowe", "4.6", "6 technik"],
            ["5. Demonstracja z omowieniem", "5", "Scenariusz 9 krokow"],
        ],
        [3600, 3300, 2460],
        header_fill=BLUE_GRAY,
    )


def main() -> None:
    diagrams = build_diagrams()
    doc = Document()
    configure_document(doc)
    add_front_matter(doc)
    section_intro(doc)
    section_spec(doc, diagrams)
    section_schedule_team(doc)
    section_architecture(doc, diagrams)
    section_demo_and_matrix(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
